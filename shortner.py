import os
import re
import sys
from docx import Document
from docx.shared import Inches, Pt

if len(sys.argv) != 2:
    sys.exit("Usage: python " + sys.argv[0] + " resume.docx")

# Test if file is a docx file.
try:
    filename = sys.argv[1]
    Document(filename)
except:
    sys.exit(filename + " is not a valid docx file.")

doc = Document(filename)

def modify_margins():
    # Change the page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(.25)
        section.bottom_margin = Inches(.25)
        section.left_margin = Inches(.25)
        section.right_margin = Inches(.25)


def modify_fonts():
    # Modify the font face and size
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)


def delete_paragraph(paragraph):
    if len(paragraph.text) > 0:
        return
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def delete_extra_spaces():
    delete_ids = []
    # From the front, find spaces
    for idx,p in enumerate(doc.paragraphs):
        if len(p.text) == 0:
            delete_ids.append(idx)
        else:
            break

    for id in delete_ids:
        if id > 0:
            p = doc.paragraphs[0]
            if len(p.text) > 0:
                break
            else:
                id = 0
        p = doc.paragraphs[id]
        delete_paragraph(p)

    # From the back, find spaces
    delete_ids = []
    num_p = len(doc.paragraphs)
    for idx,p in enumerate(reversed(doc.paragraphs)):
        num_p -= 1
        if len(p.text) == 0:
            delete_ids.append(num_p)
        else:
            break

    for id in delete_ids:
        p = doc.paragraphs[id]
        delete_paragraph(p)



def delete_overused_phrases():
    pattern = re.compile("references (available|upon)", re.IGNORECASE)
    delete_id = None

    for idx,p in enumerate(doc.paragraphs):
        if pattern.search(p.text):
            delete_id = idx

    if delete_id is not None:
        p = doc.paragraphs[delete_id]
        delete_paragraph(p)


# call functions
modify_margins()
modify_fonts()
delete_extra_spaces()
delete_overused_phrases()

# save file
filename = re.sub(r"\.docx$", "", filename, re.IGNORECASE)
doc.save(filename + '-short.docx')
