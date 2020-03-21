'''
Test for removing blank pages
'''
import os
import re
import docx
import pytest
samples = "./samples/invalid/"
tmpdir = "./samples/tmp/"

# https://github.com/python-openxml/python-docx/issues/33
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

doc = docx.Document(samples + "references.docx")
pattern = re.compile("references (available|upon)", re.IGNORECASE)
delete_id = None

for idx,p in enumerate(doc.paragraphs):
    if pattern.search(p.text):
        delete_id = idx

if delete_id is not None:
    p = doc.paragraphs[delete_id]
    delete_paragraph(p)

referencesfile = tmpdir + 'noreferences.docx'
doc.save(referencesfile)
doc = docx.Document(referencesfile)

def test_references():
    references = [p for p in doc.paragraphs if pattern.search(p.text)]
    assert len(references) == 0

try:
    os.remove(referencesfile)
except OSError as e:  ## if failed, report it back to the user ##
    print("Error: %s - %s." % (e.filename, e.strerror))
