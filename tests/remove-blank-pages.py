'''
Test for removing blank pages
'''
import os
import re
import docx
import pytest
samples = "./samples/invalid/"
tmpdir = "./samples/tmp/"

# https://github.com/python-openxml/python-docx/issues/33
def delete_paragraph(paragraph):
    if len(paragraph.text) > 0:
        return
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

doc = docx.Document(samples + "spaces.docx")
delete_ids = []

# From the front, find spaces
for idx,p in enumerate(doc.paragraphs):
    if len(p.text) == 0:
        delete_ids.append(idx)
    else:
        break

for id in delete_ids:
    if id > 0:
        p = doc.paragraphs[0]
        if len(p.text) > 0:
            break
        else:
            id = 0
    p = doc.paragraphs[id]
    delete_paragraph(p)

# From the back, find spaces
# reset
delete_ids = []
num_p = len(doc.paragraphs)
for idx,p in enumerate(reversed(doc.paragraphs)):
    num_p -= 1
    if len(p.text) == 0:
        delete_ids.append(num_p)
    else:
        break

for id in delete_ids:
    p = doc.paragraphs[id]
    delete_paragraph(p)

spacesfile = tmpdir + 'nospaces.docx'
doc.save(spacesfile)
doc = docx.Document(spacesfile)

def test_front_spaces():
    # From the front, find spaces
    delete_ids = []
    for idx,p in enumerate(doc.paragraphs):
        if len(p.text) == 0:
            delete_ids.append(idx)
        else:
            break

    assert len(delete_ids) == 0

def test_back_spaces():
    # From the back, find spaces
    delete_ids = []
    for idx,p in enumerate(reversed(doc.paragraphs)):
        if len(p.text) == 0:
            delete_ids.append(idx)
        else:
            break

    assert len(delete_ids) == 0

try:
    os.remove(spacesfile)
except OSError as e:  ## if failed, report it back to the user ##
    print("Error: %s - %s." % (e.filename, e.strerror))
